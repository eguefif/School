import pandas as pd
from docx import Document
from docx.shared import Inches
import tools as tools
import pathlib
import Comment.comment as com
#import plotly.express as px
import os


def check_name_list(name):
    data = pd.read_excel('./info.xlsx', sheet_name='names', index_col=0)
    if name in data.index:
        return True
    return False


def make_all_reports():
    return


def make_quick_summary(name, subjects, document, path):
    for subject in subjects.keys():
        levels = {'r': [], 'theta': []}

        for sub_subject in subjects[subject]:
            file = path + subject + '_' + sub_subject
            data = pd.read_excel(file, index_col=0, sheet_name=None)
            sub_subject_levels = []
            for i in data.keys():
                df = com.process_triangulation(data[i], i)
                level = round(df.loc[name, :].mean())
                sub_subject_levels.append(level)
            levels['theta'].append(subject)
            levels['r'].append(sub_subject_levels.mean())

            radar = pd.DataFrame(levels)
            fig = px.line_polar(radar, r='r', theta='theta', line_close=True)
            fig.write_image("image.png")
            document.add_picture('image.png', width=Inches(2.5))
            os.remove("image.png")
    return document


def report(name, periode):
    print('Print report for: ', name)
    if not check_name_list(name):
        print("This student's name is not in the list.")
        return

    if name == 'all':
        make_all_reports()
        return

    document = Document()
    document.add_heading(name)

    path = tools.get_parameter(periode)
    basepath = pathlib.Path(repr(path)[1:-1])
    files = basepath.iterdir()

    subjects = tools.get_subjects(files)
    document = make_quick_summary(name, subjects, document, path)
    document.save()
    return


'''
To do:
function quick_summary
This function build two radar chart for french and maths.
It makes an average for each xlx by merging every sheet.
Then it place on a radar chart.

function skills
make a paragraph for every level on the scale:
-write_extending
-proficienty
-developping
-emerging
'''

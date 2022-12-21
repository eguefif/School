import pandas as pd
import pathlib
from docx import Document
import re


def get_path(number):
    '''
    This function get the path to the required path from info.xlsx
    '''
    periode = 'P' + number
    parameters = pd.read_excel('./info.xlsx',
                               sheet_name='parameters',
                               index_col=0,)
    path = parameters.loc[periode]
    return path.values[0]


def get_subject(file, type):
    '''
    This function create a dictionnary with the name
    of all files for a specific subject
     '''
    if type == 2:
        if file.is_file():
            if re.search('_', str(file)):
                tab = str(file).rpartition('_')
                if re.search('maths', tab[0]):
                    return 0
                subject = str(file).rpartition('_')[2]
                print(subject)

    if type == 3:
        if file.is_file():
            if re.search('_', str(file)):
                tab = str(file).rpartition('_')
                if re.search('fr', tab[0]):
                    return 0
                subject = str(file).rpartition('_')[2]
                print(subject)

    if type == 1:
        if file.is_file():
            if re.search('_', str(file)):
                tab = str(file).rpartition('_')
                subject = str(file).rpartition('_')[2]
                print(subject)

    return subject


def get_match(key):
    '''
    This function get the match from the key using info.xslx
    '''
    df = pd.read_excel('./info.xlsx',
                       index_col='index',
                       sheet_name='match',)
    return df.at[key, 'Diag']


def rundown_observations_subject(data):
    '''
    Build a table for one subject with a list of required students for whom
    we don't have observations per sub_skill.
    '''
    table = {}
    data_isnan = data.isna()
    for i in range(1, data.shape[1]):
        students = []
        for j in range(data.shape[0]):
            if data_isnan.iloc[j, i]:
                students.append(data.index[j])
        if len(students) > 0:
            table[data.columns[i]] = students

    return table


def get_number_of_students():
    data = pd.read_excel('./info.xlsx', sheet_name='names', index_col=0)
    return data.shape[0]


def write_report(data, document):
    '''
    Write for each subject the missing data.
    '''
    for item in data.items():
        if len(item[1]) > 10:
            continue
        else:
            document.add_heading(item[0], level=3)
            for j in item[1]:
                document.add_paragraph(j, style='List Bullet')
    return document


def rundown_observations(periode, type):
    document = Document()
    path = get_path(periode)
    basepath = pathlib.Path(repr(path)[1:-1])
    files = basepath.iterdir()

    # Read files in order to get the subject and makes rundowns.
    for file in files:
        if re.search("_", str(file)):
            if file.is_file():
                subject = get_subject(file, type)
                if subject == 0:
                    continue
                data = pd.read_excel(file, sheet_name=None, index_col=0)
                document.add_heading(get_match(subject))

                # We have every sheets in a dictionnary, we iterate in it.
                for df in data.items():
                    rundown = rundown_observations_subject(df[1])
                    document = write_report(rundown, document)

    document.save('rundown.docx')
    print('Report saved in rundown.docx')
    return

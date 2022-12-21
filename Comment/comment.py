import re
import pandas as pd
import pathlib
from docx import Document
import tools as tools


def get_preposition(phrase):
    '''
    This function find out what to use between de and d'
    '''
    if phrase[0] in ['a', 'e', 'i', 'o', 'u', 'y']:
        return "d'"
    else:
        return 'de '


def get_gender_adj(name):
    if tools.get_gender(name) == 'm':
        return ''
    return 'e'


def write_ext(ext, pronoun, name, comment):
    if len(ext) == 1:
        comment += f"{name} a montré qu'{pronoun} avait dépassé les attentes" \
            f" pour {tools.get_match(ext[0])}. "
    if len(ext) > 1:
        comment += f"{name} a montré qu'{pronoun} avait dépassé les attentes" \
            " pour "
        for i in range(len(ext)):
            comment += f"{tools.get_match(ext[i])}"
            if i == (len(ext)-1):
                comment += f'. '
            if i == (len(ext) - 2):
                comment += ' et '
            if i < (len(ext) - 2):
                comment += ', '
    return comment


def write_proficient(proficient, pronoun, name, comment,
                     connector, commentary):
    if (re.search('était capable', commentary)
            and not re.search('était compétent', commentary)):
        if len(proficient) == 1:
            comment += f"{connector}{name} a montré qu'{pronoun} était " \
                f"compétent{get_gender_adj(name)} pour " \
                f"{tools.get_match(proficient[0])}. "
        if len(proficient) > 1:
            comment += f"{connector}{name} a montré qu'{pronoun} était " \
                f"compétent{get_gender_adj(name)} pour "
            for i in range(0, len(proficient)):
                comment += f"{tools.get_match(proficient[i])}"
                if i == (len(proficient)-1):
                    comment += f'. '
                if i == (len(proficient) - 2):
                    if re.search('compétent', comment):
                        comment += ' et '
                    else:
                        comment += ' et ' \
                            f"{get_preposition(tools.get_match(proficient[i+1]))}"
                if i < (len(proficient) - 2):
                    if re.search('compétent', comment):
                        comment += ' et '
                    else:
                        comment += ' et ' \
                            f"{get_preposition(tools.get_match(proficient[i+1]))}"
        return comment

    if len(proficient) == 1:
        comment += f"{connector}{name} a montré qu'{pronoun} était capable " \
            f"{get_preposition(tools.get_match(proficient[0]))}" \
            f"{tools.get_match(proficient[0])}. "
    if len(proficient) > 1:
        comment += f"{connector}{name} a montré qu'{pronoun} était capable " \
            f"{get_preposition(tools.get_match(proficient[0]))}"
        for i in range(0, len(proficient)):
            comment += f"{tools.get_match(proficient[i])}"
            if i == (len(proficient)-1):
                comment += f'. '
            if i == (len(proficient) - 2):
                comment += f' et {get_preposition(tools.get_match(proficient[i+1]))}'
            if i < (len(proficient) - 2):
                comment += f', {get_preposition(tools.get_match(proficient[i+1]))}'
    return comment


def write_developping(developping, pronoun, name, comment,
                      connector, commentary):
    if (re.search('encore progresser', commentary)
            and not re.search('faire des', commentary)):
        if len(developping) == 1:
            comment += f"{connector}{name} peut faire des progrès dans " \
                f"sa capacité à {tools.get_match(developping[0])}. "
        if len(developping) > 1:
            comment += f"{connector}{name} peut faire des progrès dans " \
                "sa capacité à "
            for i in range(len(developping)):
                comment += f"{tools.get_match(developping[i])}"
                if i == (len(developping)-1):
                    comment += f'. '
                if i == (len(developping) - 2):
                    comment += ' et '
                if i < (len(developping) - 2):
                    comment += ', '
        return comment

    if len(developping) == 1:
        comment += f"{connector}{name} peut encore progresser en travaillant " \
            f"sa capacité à {tools.get_match(developping[0])}. "
    if len(developping) > 1:
        comment += f"{connector}{name} peut encore progresser en travaillant " \
            "sa capacité à "
        for i in range(len(developping)):
            comment += f"{tools.get_match(developping[i])}"
            if i == (len(developping)-1):
                comment += f'. '
            if i == (len(developping) - 2):
                comment += ' et '
            if i < (len(developping) - 2):
                comment += ', '
    return comment


def write_emerging(emerging, pronoun, name, comment, connector):
    if len(emerging) == 1:
        comment += f"{name} a {connector}une marge importante de progression " \
            f"pour {tools.get_match(emerging[0])}. "
    if len(emerging) > 1:
        comment += f"{name} a {connector}une marge importante de" \
            " progression pour "
        for i in range(len(emerging)):
            comment += f"{tools.get_match(emerging[i])}"
            if i == (len(emerging)-1):
                comment += f'. '
            if i == (len(emerging) - 2):
                comment += ' et '
            if i < (len(emerging) - 2):
                comment += ', '
    return comment


def sub_subject_comment(name, data, subject, commentary):
    '''
This function built a commentary for a sub skill such as ecriture
    '''
    gender = tools.get_gender(name)
    if str(gender) == 'm':
        pronoun = 'il'
        c_pronoun = 'Il'
    else:
        pronoun = 'elle'
        c_pronoun = 'Elle'
    proficient = []
    ext = []
    developping = []
    emerging = []
    connector = ''
    comment = tools.get_match(subject) + ', '

    # We create a list with all different level of proficienty

    values = data.loc[name].iloc[:].values
    for j in range(len(values)):
        if values[j] == 1:
            emerging.append(data.columns[j])
        if values[j] == 2:
            developping.append(data.columns[j])
        if values[j] == 3:
            proficient.append(data.columns[j])
        if values[j] > 3:
            ext.append(data.columns[j])

    comment = write_ext(ext, pronoun, name, comment,)

    # Now editing proficient
    if len(ext) > 0:
        connector = ' De plus, '
        name = pronoun

    comment = write_proficient(proficient,
                               pronoun,
                               name,
                               comment,
                               connector,
                               commentary)

    if len(ext) > 0 or len(proficient) > 0:
        connector = ' Maintenant, '
        name = pronoun

    comment = write_developping(developping,
                                pronoun,
                                name,
                                comment,
                                connector,
                                commentary)

    connector = ''
    name = c_pronoun

    if len(ext) > 0 or len(proficient) > 0:
        if len(developping) == 0:
            connector = 'cependant '
            name = c_pronoun
    if len(developping) > 0:
        connector = 'aussi '
        name = c_pronoun

    comment = write_emerging(emerging,
                             pronoun,
                             name,
                             comment,
                             connector,)

    comment += '\n'
    return comment


def subject_comment(data, doc, subject):
    ''''Function that build comments for one subject such as fr or maths.'''
    doc.add_heading(tools.get_match(subject))
    list_names = tools.get_list_names()

    for name in list_names:
        comment = name + '\n'
        for sub_subject in data.keys():
            comment += sub_subject_comment(name,
                                           data[sub_subject],
                                           sub_subject,
                                           comment)
            continue
        doc.add_paragraph(comment)
        continue

    return doc


def process_triangulation(data, sub_subject):
    # Most of the time, the teach would have get observation by the time he or
    # she is making the commentaries. If not, we put a 2 and print a warning
    # saying there was a fillna
    if data.isnull().values.any():
        print(f'We are filling missing data with 2 in {sub_subject}')
        data.fillna(2, inplace=True)

    # replacing triangulation by a score
    threshold = tools.get_parameter('threshold')
    for i in range(data.shape[0]):
        for j in range(data.shape[1]):
            # if there is string, means it is a triangulation
            if isinstance(data.iloc[i, j], str):
                if len(data.iloc[i, j]) >= threshold:
                    # Lowercase mean proficient 3, uppercase means ext 4
                    if data.iloc[i, j].islower():
                        data.iloc[i, j] = 3
                    else:
                        data.iloc[i, j] = 4
                else:
                    data.iloc[i, j] = 2
    return data


def get_average(data, sub_subject):
    '''We build a summary with averaged for each sheet in the spreadsheet and
    for each student
    '''
    average = pd.DataFrame(index=tools.get_list_names())
    for skill in data:
        notes = []
        for i in range(len(data[skill])):
            if data[skill].iloc[i, 1:].isnull().sum() == data[skill].shape[1]:
                continue
            sk = process_triangulation(data[skill], sub_subject)

            notes.append(
                tools.round_half(sk.iloc[i, 0:].values.sum()/sk.shape[1]))

        average[skill] = notes

    return average


def make_comments(subjects, path):
    ''' This function create a docx with comments for all subjects'''
    doc = Document()

    # For each subject, we build a new dataframe
    for subject in subjects.keys():
        df = {}
        for sub_subject in subjects[subject]:
            file = path + subject + '_' + sub_subject
            data = pd.read_excel(
                file, sheet_name=None, index_col=0)
            df[sub_subject] = get_average(data, sub_subject)

        # We call the function to built a comment for a subject
        doc = subject_comment(df, doc, subject)

    return doc


def comments(periode):
    '''
Core function of the comment app.
Periode : this parameter target the periode(term) of the year wished.
p1 for periode 1
p2 for periode 2
p3 for periode 3
    '''

    path = tools.get_parameter(periode)
    basepath = pathlib.Path(repr(path)[1:-1])
    files = basepath.iterdir()

    subjects = tools.get_subjects(files)
    doc_comment = make_comments(subjects, path)
    doc_comment.save('./comment.docx')
    print('Comment saved in comment.docx')


'''
To do:
Function that process each spreadsheet by:
 -solve bugs
'''
